import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Fehler: Die Bibliothek 'python-docx' ist nicht installiert.")
    print("Bitte installieren Sie sie mit: pip install python-docx")
    sys.exit(1)

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = paragraph._p.make_element("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id, )

    new_run = paragraph._p.make_element("w:r")
    rPr = new_run.make_element("w:rPr")

    if color:
        c = new_run.make_element("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = new_run.make_element("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    text_elem = new_run.make_element("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_and_add_runs(paragraph, text):
    # Regex für Code (`...`), Links ([...](...)) und Bold (**...**)
    # Wir splitten den Text und behalten die Trenner
    # Pattern erklärt:
    # (`[^`]+`) -> Code Blöcke
    # (\[[^\]]+\]\([^)]+\)) -> Links
    # (\*\*[^*]+\*\*) -> Bold Text
    pattern = r'(`[^`]+`|\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        elif part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            # Link
            try:
                # Suche nach [text](url)
                m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                if m:
                    link_text = m.group(1)
                    link_url = m.group(2)
                    add_hyperlink(paragraph, link_url, link_text)
                else:
                    paragraph.add_run(part)
            except:
                paragraph.add_run(part)
        else:
            # Normaler Text
            paragraph.add_run(part)

def markdown_to_docx(md_file, docx_file):
    print(f"Konvertiere {md_file} nach {docx_file}...")
    
    if not os.path.exists(md_file):
        print(f"Datei {md_file} nicht gefunden.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Styles anpassen
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    in_code_block = False
    code_block_content = []

    for line in lines:
        line = line.rstrip()
        
        # Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # Ende des Code Blocks
                p = doc.add_paragraph()
                text = '\n'.join(code_block_content)
                runner = p.add_run(text)
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                p.style = 'No Spacing'
                code_block_content = []
                in_code_block = False
            else:
                # Start des Code Blocks
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            parse_and_add_runs(p, line[2:])
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            parse_and_add_runs(p, line[3:])
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            parse_and_add_runs(p, line[4:])
        elif line.startswith('#### '):
            p = doc.add_heading(level=4)
            parse_and_add_runs(p, line[5:])
        
        # Horizontal Rule
        elif line.startswith('---') or line.startswith('***'):
            p = doc.add_paragraph()
            run = p.add_run('_' * 80)
            run.font.color.rgb = RGBColor(200, 200, 200)
            
        # Bullet Points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_and_add_runs(p, text)
            
        # Numbered Lists
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            parse_and_add_runs(p, text)
            
        # Empty lines
        elif line.strip() == '':
            continue
            
        # Normal Text
        else:
            p = doc.add_paragraph()
            parse_and_add_runs(p, line)

    doc.save(docx_file)
    print(f"✅ {docx_file} erfolgreich erstellt!")

if __name__ == "__main__":
    markdown_to_docx("README.md", "Installationsanleitung.docx")
    markdown_to_docx("QUICKSTART.md", "Schnellanleitung.docx")
    markdown_to_docx("USER_MANUAL.md", "Anwenderhandbuch.docx")
